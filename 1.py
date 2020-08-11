from docx import Document
import os
import pandas as pd


NUMBER_GROUP = input('Введите номер группы:  ')
directory = input('Вставьте путь к папке с файлами: ')
directory_save = input('Вставьте путь к папке в которую сохранить конечный файл: ')
TITLE = ['ФИО', 'Вид поощреения', 'Формулировка', 'Кем поощрен']

###Собираем таблицу из всех файлов в папке ###

def table_of_all_files_from_the_directory(doc):###передаем в функцию список документов в папке (функция list_docx)
    finish_table =[] #Конечная таблица из файлов
    for file in doc: # для каждого документа
        file_directory = directory + '\\' + file # для каждого документа склеили путь
        DOCUMENT = Document(file_directory) # открываем документ, присваиваем его переменной
        final_table = table_names_promoyion(DOCUMENT) #запускаем функцию которая таблицу с ФИО поощрением, передаем ей переменную с документом
        for string in final_table: # каждую строку из получившейся таблицы
            finish_table.append(string)# сохраняем в конечную таблицу
    return (finish_table)




#Создаем и возвращаем список документов в каталоге
# функция на вход принимаем переменную с путем к папке
def list_docx (directory):
    #Каталог из которого будем брать изображения
    directory = directory
    #Получаем список файлов в переменную files
    files = os.listdir(directory)
    #Фильтруем список по расширению файлов
    docx = filter(lambda x: x.endswith('.docx'), files)
    return (docx)



### Функция возвращает таблицу с ФИО поощрением и тем кто поощрил ###
# функция на вход принимает переменную с файлом для обработки
def table_names_promoyion(doc):
    table1 = doc.tables[0]  # берем из документа первую таблицу (индекс 0)
    list_columns = table1.columns  # получаем список столбцов
    names_promotion = []  # конечная таблица ФИО и поощрение
    for number_element in range(
            len(table1.rows)):  # от 0 до последней строки len(table1.rows) - получили количество строк в таблице
        h = []  # промежуточный список в который собираем строку таблицы

        ### берем каждую ячейку из строки, обрабатываем и добавляем в список h  ###
        for z in list_columns:  #  беремстроку z из list_columns , подставляем в нее индекс ячейки number_element
            h.append(z.cells[number_element].text.strip().replace('\n',
                                                                  ' '))  # получаем ячейку .text - преобразуем в текст # .strip()
            # - убираем лишние пробелы в начале и конце строки
            # .replace('\n', ' ') - заменяем символ \n на пробел, добавляем во временный список h

        if h[5] == NUMBER_GROUP:  # проверяем условие, если номер группы в ячейке равен номеру нужной группы
            n = h[6].replace('-', '–').split('–')  # заменяем символ тире на символ дефис, разбиваем ячейку с поощрением по разделителю дефис (отделяем поощрение от формулировки)
            a = [h[2], n[0], n[1], h[1]]  # создаем временный список в который помещаем ячейку с ФИО и благодарностью
            names_promotion.append(a)  # добавляем временный список в конечную таблицу
    return (names_promotion)

######Собираем DataFrame ######

gg = table_of_all_files_from_the_directory(list_docx (directory)) #Присваиваем переменной готовую таблицу
dm = pd.DataFrame(data = gg, columns = TITLE) # Создаем DtaFrame
hf = dm.groupby(['ФИО','Вид поощреения']).size().unstack() #групппируем таблицу по ФИО и поощрениям, получаем количество каждоого поощрения


   ###Сохраняем новую таблицус поощрениями в новый документ###

def new_list_group(doc, df, name): # функция на вход принимает конечную таблицу, DataFrame с количеством поощрений  и название нового файла
    Name_doc = name +'.docx' # склеиваем названме файла
    directory_finish = directory_save + '\\' + Name_doc # директория куда сохраняем файл
    docfinish = Document() #Создаем пустой документ
    tables = doc # сохраняем в переменную готовую таблицу
    table = docfinish.add_table(rows = len(tables), cols = len(tables[0])) # Создаем в документи таблицу размером:
    table.style = 'Table Grid' #Стиль таблицы
    # кол во строк == кол ву строк таблицы tables, кол во столбцов == кол ву столбцов таблицы tables
# создаем цикл для записи в документ
    for row in range(len(tables)): # ,берем каждую строку в диапазоне длинны таблицы tables
        for b in range(len(tables[0])): #берем каждый столбец в диапазоне длинны первой строки таблицы tables
            cell = table.cell(row, b) # выбираем ячейку с индексом row, b
            cell.text = tables[row][b] #записываем в ячейку данные из ячейки таблицы tables с индексом row, b

    docfinish.add_paragraph('Общее число поощрений') # Добавляем параграф с названием новой таблицы
### Добавляем вторую таблицу с общим количеством поощрений ###

    tables_1 = df # присваиваем переменной таблицу DataFrame

    table1 = docfinish.add_table(tables_1.shape[0] + 2, tables_1.shape[1] + 1) #Добавляем вторую таблицу
    table1.style = 'Table Grid'  # Стиль таблицы
    for j in range(tables_1.shape[-1]):
        table1.cell(0, j + 1).text = tables_1.columns[j]

    for i in range(tables_1.shape[0]):
        for j in range(tables_1.shape[-1]):
            table1.cell(i + 2, j + 1).text = str(tables_1.values[i, j])

    # написать название столбцов в файл
    table1.cell(0, 0).text = tables_1.columns.name

    # писать значения в файл
    table1.cell(1, 0).text = tables_1.index.name

    for i, my_index in enumerate(tables_1.index):
        table1.cell(i + 2, 0).text = my_index
    docfinish.save(directory_finish) # сохраняем документ
new_list_group(table_of_all_files_from_the_directory(list_docx (directory)), hf, 'Список поощрений за группу')















